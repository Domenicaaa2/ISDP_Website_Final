import io
import os
import re
import time
import asyncio
import logging
from datetime import datetime
from typing import Optional

import httpx
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)

app = FastAPI(title="ISDP Platform")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

IBM_API_KEY    = os.getenv("IBM_API_KEY")
WXO_URL        = os.getenv("WXO_URL", "https://api.eu-de.watson-orchestrate.cloud.ibm.com")
ENVIRONMENT_ID = os.getenv("ENVIRONMENT_ID", "draft")

# ── IAM token cache ───────────────────────────────────────────────────────────
_token_cache: dict = {"token": None, "expires_at": 0.0}


async def get_token() -> str:
    if _token_cache["token"] and time.time() < _token_cache["expires_at"] - 60:
        return _token_cache["token"]
    if not IBM_API_KEY:
        raise HTTPException(500, "IBM_API_KEY not set")
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.post(
            "https://iam.cloud.ibm.com/identity/token",
            data={"grant_type": "urn:ibm:params:oauth:grant-type:apikey", "apikey": IBM_API_KEY},
            headers={"Content-Type": "application/x-www-form-urlencoded"},
        )
        if not r.is_success:
            raise HTTPException(401, f"IAM token exchange failed: {r.text[:300]}")
        d = r.json()
        _token_cache["token"] = d["access_token"]
        _token_cache["expires_at"] = time.time() + d["expires_in"]
        logger.info("IAM token refreshed (expires in %ss)", d["expires_in"])
        return _token_cache["token"]


# ── Models ────────────────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    agent_id: str
    thread_id: Optional[str] = None


# ── File upload & text extraction ─────────────────────────────────────────────

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    content = await file.read()
    filename = file.filename or "document"
    text = ""

    try:
        if filename.lower().endswith(".pdf"):
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))

            # Build page-index → chapter title map from PDF bookmarks/outline
            chapter_map: dict = {}
            try:
                def _walk_outline(items: list, depth: int = 0) -> None:
                    for item in items:
                        if isinstance(item, list):
                            _walk_outline(item, depth + 1)
                        elif hasattr(item, "title"):
                            try:
                                page_idx = reader.get_destination_page_number(item)
                                chapter_map.setdefault(page_idx, item.title)
                            except Exception:
                                pass
                _walk_outline(reader.outline)
            except Exception:
                pass

            # Extract text page by page, inserting [Kapitel: ...] markers when
            # the bookmark map says a new chapter starts on this page.
            chunks: list = []
            for page_idx, page in enumerate(reader.pages):
                if page_idx in chapter_map:
                    chunks.append(f"[Kapitel: {chapter_map[page_idx]}]")
                page_text = page.extract_text()
                if page_text:
                    chunks.append(page_text)
            text = "\n\n".join(chunks)
        elif filename.lower().endswith((".docx", ".doc")):
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif filename.lower().endswith((".txt", ".md")):
            text = content.decode("utf-8", errors="replace")
        else:
            text = content.decode("utf-8", errors="replace")
    except Exception as e:
        logger.warning("Could not extract text from %s: %s", filename, e)
        text = f"[Could not extract text: {e}]"

    if len(text) > 8000:
        text = text[:8000] + "\n\n[... document truncated ...]"

    logger.info("Uploaded %s (%d chars extracted)", filename, len(text))
    return {"filename": filename, "text": text, "chars": len(text)}


# ── WatsonX Orchestrate run / poll ────────────────────────────────────────────

FLOW_STUB_MARKERS = (
    "a new flow has started",
    "this chat session is currently dedicated to the flow",
    "will resume once the flow is complete",
)


def _looks_like_flow_stub(text: Optional[str]) -> bool:
    if not text:
        return False
    t = text.strip().lower()
    return any(m in t for m in FLOW_STUB_MARKERS)


def _extract_flow_candidates(payload, stub_text: str) -> list:
    """Extract texts that look like real flow output from a WXO payload.

    Handles two known shapes:
      - /threads/{tid}/messages → list of {role, content:[{text, ...}]}
      - /runs?thread_id={tid}  → {data:[{result:{data:{message:{content:[{text}]}}}}]}

    Filters out user messages, stub messages, and tiny snippets.
    """
    stub_norm = (stub_text or "").strip().lower()
    out: list = []

    def keep(t: Optional[str]) -> Optional[str]:
        if not t or not isinstance(t, str):
            return None
        s = t.strip()
        if len(s) < 60:
            return None
        if _looks_like_flow_stub(s):
            return None
        if s.lower() == stub_norm:
            return None
        return s

    if isinstance(payload, list):
        for msg in payload:
            if not isinstance(msg, dict) or msg.get("role") != "assistant":
                continue
            for item in msg.get("content") or []:
                if isinstance(item, dict):
                    kept = keep(item.get("text"))
                    if kept:
                        out.append(kept)
    elif isinstance(payload, dict):
        # runs list
        for run in payload.get("data") or []:
            if not isinstance(run, dict):
                continue
            try:
                content = (
                    ((run.get("result") or {}).get("data") or {})
                    .get("message", {})
                    .get("content")
                    or []
                )
            except AttributeError:
                content = []
            for item in content:
                if isinstance(item, dict):
                    kept = keep(item.get("text"))
                    if kept:
                        out.append(kept)
    return out


async def _probe_thread_endpoints(
    client: httpx.AsyncClient, headers: dict, thread_id: str
) -> list:
    """Find which WXO thread-related GET endpoints are accepted by the
    current instance. Returns the list of working absolute URLs."""
    candidates = [
        f"{WXO_URL}/v1/orchestrate/threads/{thread_id}/messages",
        f"{WXO_URL}/v1/orchestrate/threads/{thread_id}/runs",
        f"{WXO_URL}/v1/orchestrate/threads/{thread_id}",
        f"{WXO_URL}/v1/orchestrate/threads/{thread_id}/events",
        f"{WXO_URL}/v1/orchestrate/runs?thread_id={thread_id}",
    ]
    working = []
    for url in candidates:
        try:
            r = await client.get(url, headers=headers, timeout=10)
            if r.is_success:
                working.append(url)
                logger.info("Thread endpoint OK: %s", url.replace(WXO_URL, ""))
            else:
                logger.debug("Thread endpoint %s -> %d", url.replace(WXO_URL, ""), r.status_code)
        except Exception as e:
            logger.debug("Thread endpoint error %s: %s", url, e)
    return working


async def _wait_for_flow_output(
    client: httpx.AsyncClient,
    headers: dict,
    thread_id: str,
    stub_text: str,
    max_seconds: int = 240,
) -> Optional[str]:
    """After a flow agent returns the stub, poll the thread until the
    real flow output appears."""
    working = await _probe_thread_endpoints(client, headers, thread_id)
    if not working:
        logger.warning("No thread endpoint responded — cannot wait for flow output.")
        return None

    stub_norm = stub_text.strip().lower()
    deadline = time.time() + max_seconds
    seen_signatures: set = set()
    poll_count = 0

    while time.time() < deadline:
        await asyncio.sleep(3)
        poll_count += 1
        for url in working:
            try:
                r = await client.get(url, headers=headers, timeout=15)
                if not r.is_success:
                    continue
                try:
                    payload = r.json()
                except Exception:
                    continue
                interesting = _extract_flow_candidates(payload, stub_text)
                if interesting:
                    # Pick the longest — most likely the actual draft
                    interesting.sort(key=len, reverse=True)
                    winner = interesting[0]
                    logger.info(
                        "Flow output received via %s after %d polls (%d chars).",
                        url.replace(WXO_URL, ""), poll_count, len(winner),
                    )
                    return winner
                seen_signatures.add((url, poll_count))
            except Exception as e:
                logger.debug("Flow polling error on %s: %s", url, e)

    logger.warning(
        "Flow output not received after %ds. Seen %d distinct text signatures.",
        max_seconds, len(seen_signatures),
    )
    return None


async def _start_run(
    client: httpx.AsyncClient,
    headers: dict,
    agent_id: str,
    message: str,
    thread_id: Optional[str],
    environment_id: Optional[str],
) -> dict:
    body: dict = {
        "agent_id": agent_id,
        "message": {"role": "user", "content": message},
    }
    if environment_id:
        body["environment_id"] = environment_id
    if thread_id:
        body["thread_id"] = thread_id
    r = await client.post(f"{WXO_URL}/v1/orchestrate/runs", headers=headers, json=body)
    return r


async def run_wxo(
    client: httpx.AsyncClient,
    headers: dict,
    agent_id: str,
    message: str,
    thread_id: Optional[str],
) -> tuple[str, str]:
    # Try with configured environment_id first; fall back to no environment_id
    # if WXO says the agent is not found in that environment.
    r = await _start_run(client, headers, agent_id, message, thread_id, ENVIRONMENT_ID)

    if not r.is_success:
        err_text = r.text
        logger.warning("Run with env=%s failed %s: %s", ENVIRONMENT_ID, r.status_code, err_text[:300])
        # If the error mentions environment, retry without environment_id
        if "environment" in err_text.lower() or r.status_code == 500:
            logger.info("Retrying without environment_id for agent %s", agent_id)
            r = await _start_run(client, headers, agent_id, message, thread_id, None)
        if not r.is_success:
            logger.error("Run failed %s: %s", r.status_code, r.text[:800])
            raise HTTPException(r.status_code, f"Run start failed: {r.text[:800]}")

    run_data = r.json()
    current_run_id: str = run_data["run_id"]
    returned_thread_id: str = run_data["thread_id"]
    accumulated_text = ""

    logger.info("Run started: %s  thread: %s", current_run_id, returned_thread_id)

    for i in range(200):
        await asyncio.sleep(1)
        poll = await client.get(
            f"{WXO_URL}/v1/orchestrate/runs/{current_run_id}", headers=headers
        )
        if not poll.is_success:
            raise HTTPException(poll.status_code, f"Poll failed: {poll.text[:200]}")

        result = poll.json()
        status: str = result["status"]
        logger.info("Poll %d (run %s) status: %s", i, current_run_id, status)

        if status == "failed":
            raise HTTPException(500, f"Run failed: {result.get('last_error', 'unknown')}")

        if status == "completed":
            content = result["result"]["data"]["message"]["content"]
            chunk = " ".join(c.get("text", "") for c in content if isinstance(c, dict))
            accumulated_text += chunk
            next_run_id: Optional[str] = result["result"].get("next_run_id")
            if next_run_id:
                current_run_id = next_run_id
                logger.info("Continuing with next_run_id: %s", next_run_id)
            else:
                # Flow agent detection: the initial run sometimes returns just
                # a "A new flow has started…" stub while the real output is
                # produced asynchronously and posted to the thread.
                if _looks_like_flow_stub(accumulated_text):
                    logger.info(
                        "Detected flow stub from agent %s — waiting for flow output…",
                        agent_id,
                    )
                    flow_output = await _wait_for_flow_output(
                        client, headers, returned_thread_id, accumulated_text
                    )
                    if flow_output:
                        return returned_thread_id, flow_output
                    logger.warning(
                        "Flow output never arrived; returning stub to caller."
                    )
                return returned_thread_id, accumulated_text

    raise HTTPException(504, "Run timed out after 120 seconds")


# ── API endpoints ─────────────────────────────────────────────────────────────

@app.post("/api/chat")
async def chat(req: ChatRequest):
    token = await get_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=httpx.Timeout(60.0, connect=10.0)) as client:
        thread_id, response_text = await run_wxo(
            client, headers, req.agent_id, req.message, req.thread_id
        )
    return {"response": response_text, "thread_id": thread_id, "agent_id": req.agent_id}


@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "wxo_url": WXO_URL,
        "environment_id": ENVIRONMENT_ID,
        "api_key_set": bool(IBM_API_KEY),
    }


@app.get("/api/debug/thread/{thread_id}")
async def debug_thread(thread_id: str):
    """Probe all known WXO thread endpoints and return whichever respond.
    Used to figure out where a flow-agent's actual output is delivered."""
    token = await get_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    candidates = [
        f"/v1/orchestrate/threads/{thread_id}/messages",
        f"/v1/orchestrate/threads/{thread_id}/runs",
        f"/v1/orchestrate/threads/{thread_id}",
        f"/v1/orchestrate/threads/{thread_id}/events",
        f"/v1/orchestrate/runs?thread_id={thread_id}",
    ]
    result = {}
    async with httpx.AsyncClient(timeout=30) as client:
        for path in candidates:
            try:
                r = await client.get(f"{WXO_URL}{path}", headers=headers)
                result[path] = {
                    "status": r.status_code,
                    "body": r.json() if r.is_success and r.headers.get("content-type", "").startswith("application/json") else r.text[:500],
                }
            except Exception as e:
                result[path] = {"status": "error", "body": str(e)}
    return result


@app.get("/api/debug/agents")
async def debug_agents():
    """List all agents available in the WXO instance so we can verify agent IDs."""
    token = await get_token()
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    async with httpx.AsyncClient(timeout=30) as client:
        r = await client.get(f"{WXO_URL}/v1/agents", headers=headers)
        if r.is_success:
            return {"source": "/v1/agents", "data": r.json()}
        r2 = await client.get(f"{WXO_URL}/v1/orchestrate/agents", headers=headers)
        if r2.is_success:
            return {"source": "/v1/orchestrate/agents", "data": r2.json()}
        return {
            "error": f"Could not list agents.",
            "body1": r.text[:500],
            "body2": r2.text[:500],
        }


# ── ISDP Word Draft Generation ────────────────────────────────────────────────

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "ISDS_Vorlage.docx")

# Canonical template chapter headings (in document order). Used both to instruct
# the German Draft agent and to fill the template chapter-by-chapter.
ISDS_SECTIONS = [
    ("Zweck", 1),
    ("Ausgangslage / Situationsbeschreibung", 1),
    ("Sicherheitsbedarf gemäss Schutzbedarfsanalyse", 1),
    ("Datenschutz / Datenschutz-Folgenabschätzung (DSFA)", 1),
    ("Rechtsgrundlagen", 2),
    ("Bearbeitung von Personendaten", 2),
    ("Zweck der Datenbearbeitung & Aufbewahrungsdauer", 2),
    ("Personendatenkategorien", 2),
    ("Art der Bearbeitung", 2),
    ("Verarbeitungsorte", 2),
    ("Subunternehmer", 2),
    ("Sicherheitsrelevante Systembeschreibung", 1),
    ("Ansprechpersonen / Verantwortlichkeiten", 2),
    ("Beschreibung des Gesamtsystems", 2),
    ("Backup", 3),
    ("Nachvollziehbarkeit / Logging", 3),
    ("Störungsbehebung", 3),
    ("Beschreibung der zugrundeliegenden Technik / Technologien", 2),
    ("Architekturskizze / Kommunikationsmatrix", 3),
    ("Benutzer-, Rollen- und Berechtigungskonzept", 3),
    ("Authentifizierung", 3),
    ("Remote-Zugriff durch Dritte", 3),
    ("Bekannte Schwachstellen", 3),
    ("Aufbewahrung und Archivierung", 1),
    ("Notfallkonzept", 1),
    ("Risikoanalyse und Schutzmassnahmen", 1),
    ("Identifizierte Risiken", 2),
    ("Schutzmassnahmen", 2),
    ("Restrisiken", 2),
    ("Risikoübersicht (aktuelle Risiken und Restrisiken)", 2),
]

MISSING_MARKER = "_Keine Informationen aus den hochgeladenen Dokumenten verfügbar._"


class DraftRequest(BaseModel):
    project_name: Optional[str] = None
    draft_text: str


@app.get("/api/isds-sections")
async def isds_sections():
    """Expose the canonical template chapter list so the frontend can include
    it in the German Draft agent prompt."""
    return {"sections": [{"name": n, "level": l} for n, l in ISDS_SECTIONS]}


def _normalize_heading(text: str) -> str:
    """Aggressive normalization for fuzzy heading matching."""
    t = text.strip().lower()
    # strip "Kap. X.X – " prefix if agent uses old format
    t = re.sub(r"^kap\.?\s*[\d\.]+\s*[–\-]\s*", "", t)
    # collapse non-alphanumeric runs to single space
    t = re.sub(r"[^a-zäöüß0-9]+", " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


def _extract_heading_candidate(line: str) -> Optional[str]:
    """Return the raw heading text if line looks like a heading, else None.

    Recognizes:
      ## Headingtext            (any level of #)
      **Headingtext**            (bold-wrapped on its own line)
      **Kap. X.X – Headingtext** (legacy chapter format)
      1. Headingtext             (numbered, if short)
      Headingtext:               (followed by colon)
      Headingtext                (plain line that matches a known template heading)
    """
    stripped = line.strip()
    if not stripped or len(stripped) > 200:
        return None

    m = re.match(r"^#{1,6}\s+(.+?)\s*$", stripped)
    if m:
        return m.group(1).strip()

    m = re.match(r"^\*\*\s*(.+?)\s*\*\*\s*:?\s*$", stripped)
    if m:
        return m.group(1).strip()

    # Numbered-list-style heading: "1. Zweck"  or "1.1 Zweck"
    m = re.match(r"^\d+(?:\.\d+)*\.?\s+([A-ZÄÖÜ].{2,150})$", stripped)
    if m:
        return m.group(1).strip()

    # "Heading:" style
    m = re.match(r"^([A-ZÄÖÜ][^:]{3,150}):\s*$", stripped)
    if m:
        return m.group(1).strip()

    # Fallback: plain line — accept only if it strongly matches a known
    # template heading (handled by caller, returns the line as-is here).
    if 3 <= len(stripped) <= 150 and not stripped.endswith(("."  , "!", "?")):
        return stripped

    return None


# Precomputed normalized template heading set for fuzzy matching
_TEMPLATE_HEADING_NORMS = None


def _template_heading_norms():
    global _TEMPLATE_HEADING_NORMS
    if _TEMPLATE_HEADING_NORMS is None:
        _TEMPLATE_HEADING_NORMS = {_normalize_heading(n): n for n, _ in ISDS_SECTIONS}
    return _TEMPLATE_HEADING_NORMS


def _match_to_template(candidate_norm: str) -> Optional[str]:
    """Try to map a candidate heading (already normalized) to a template
    heading key. Exact match preferred, then substring containment."""
    norms = _template_heading_norms()
    if candidate_norm in norms:
        return candidate_norm
    # Substring match: prefer the longest matching template heading that is
    # either contained in the candidate or contains the candidate.
    best = None
    best_len = 0
    for tnorm in norms:
        if len(tnorm) < 5:
            continue
        if tnorm in candidate_norm or candidate_norm in tnorm:
            common = min(len(tnorm), len(candidate_norm))
            ratio = common / max(len(tnorm), len(candidate_norm))
            if ratio >= 0.6 and common > best_len:
                best = tnorm
                best_len = common
    return best


def _parse_agent_sections(text: str) -> dict:
    """Parse the German Draft agent output into {normalized_heading: content}.

    Two-pass strategy:
      1. Strict pass: lines that look like markdown headings start a section.
      2. Template-aware fallback: any line that matches a known template
         heading name (after normalization) also starts a section.

    The fallback ensures the parser works even if the agent ignores the
    requested `## Heading` format.
    """
    sections: dict = {}
    current_heading_norm: Optional[str] = None
    current_buf: list = []
    template_norms = _template_heading_norms()

    def commit():
        nonlocal current_heading_norm, current_buf
        if current_heading_norm is None:
            return
        content = "\n".join(current_buf).strip()
        # If the same heading appears twice, keep the longer content
        if current_heading_norm in sections and len(sections[current_heading_norm]) >= len(content):
            return
        sections[current_heading_norm] = content

    lines = text.split("\n")
    for raw_line in lines:
        stripped = raw_line.strip()
        accepted_norm = None

        # Only "##  Heading" (markdown heading) is treated as an unconditional
        # section start. Everything else (**bold**, numbered prefix, plain
        # line) must map to a known template heading — otherwise inline
        # emphases like "**Authentifizierung**" inside a body paragraph
        # would be mis-parsed as a new section and swallow the previous
        # section's content.
        m_hash = re.match(r"^#{1,6}\s+(.+?)\s*$", stripped)
        if m_hash:
            cand_norm = _normalize_heading(m_hash.group(1))
            if cand_norm:
                accepted_norm = _match_to_template(cand_norm) or cand_norm
        else:
            candidate = _extract_heading_candidate(raw_line)
            if candidate is not None:
                cand_norm = _normalize_heading(candidate)
                if cand_norm:
                    mapped = _match_to_template(cand_norm)
                    if mapped is not None:
                        accepted_norm = mapped

        if accepted_norm:
            commit()
            current_heading_norm = accepted_norm
            current_buf = []
            continue

        if current_heading_norm is not None:
            current_buf.append(raw_line)

    commit()
    return sections


def _append_markdown_to_doc(doc, text: str) -> None:
    """Parse a markdown-ish string (headings, lists, tables, paragraphs)
    from the German Draft agent output and append it to the docx."""

    available_styles = {s.name for s in doc.styles}

    def pick_style(*candidates):
        for c in candidates:
            if c in available_styles:
                return c
        return None

    bullet_style = pick_style("List Bullet", "List Paragraph")
    number_style = pick_style("List Number", "List Paragraph")
    table_style = pick_style("Light Grid Accent 1", "Table Grid", "Light Grid")

    def add_styled_paragraph(text_, style_name):
        p = doc.add_paragraph()
        _render_inline(p, text_)
        if style_name:
            try:
                p.style = doc.styles[style_name]
            except (KeyError, ValueError):
                pass
        return p

    _INLINE_PATTERN = re.compile(
        r"(\*\*[^*\n]+?\*\*"           # **bold**
        r"|__[^_\n]+?__"               # __bold__
        r"|\*[^*\n]+?\*"               # *italic*
        r"|_[^_\n]+?_"                 # _italic_
        r"|`[^`\n]+?`"                 # `code`
        r")"
    )

    def _render_inline(paragraph, text_):
        """Render markdown-ish inline formatting (bold, italic, code) into
        the given paragraph as styled runs."""
        if not text_:
            return
        parts = _INLINE_PATTERN.split(text_)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**") and len(part) >= 4:
                r = paragraph.add_run(part[2:-2])
                r.bold = True
            elif part.startswith("__") and part.endswith("__") and len(part) >= 4:
                r = paragraph.add_run(part[2:-2])
                r.bold = True
            elif part.startswith("*") and part.endswith("*") and len(part) >= 3:
                r = paragraph.add_run(part[1:-1])
                r.italic = True
            elif part.startswith("_") and part.endswith("_") and len(part) >= 3:
                r = paragraph.add_run(part[1:-1])
                r.italic = True
            elif part.startswith("`") and part.endswith("`") and len(part) >= 3:
                r = paragraph.add_run(part[1:-1])
                r.font.name = "Consolas"
            else:
                paragraph.add_run(part)

    lines = text.split("\n")
    i = 0
    n = len(lines)

    def flush_table(table_lines):
        rows = []
        for line in table_lines:
            if re.match(r"^\s*\|?[\s\-:|]+\|?\s*$", line):
                continue
            stripped = line.strip()
            if stripped.startswith("|"):
                stripped = stripped[1:]
            if stripped.endswith("|"):
                stripped = stripped[:-1]
            cells = [c.strip() for c in stripped.split("|")]
            rows.append(cells)
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=col_count)
        if table_style:
            try:
                table.style = doc.styles[table_style]
            except (KeyError, ValueError):
                pass
        for r_idx, row in enumerate(rows):
            for c_idx in range(col_count):
                cell = table.rows[r_idx].cells[c_idx]
                cell_text = row[c_idx] if c_idx < len(row) else ""
                # Clear the default empty paragraph and render with inline formatting
                cell.text = ""
                p = cell.paragraphs[0]
                _render_inline(p, cell_text)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            flush_table(table_lines)
            continue

        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        m = re.match(r"^\*\*Kap\.?\s*([\d\.]+)\s*[–-]\s*(.+?)\*\*\s*$", stripped)
        if m:
            doc.add_heading(f"Kap. {m.group(1)} – {m.group(2).strip()}", level=2)
            i += 1
            continue

        m = re.match(r"^\*\*(.+?)\*\*\s*$", stripped)
        if m:
            p = doc.add_paragraph()
            run = p.add_run(m.group(1).strip())
            run.bold = True
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-_*]{3,}\s*$", stripped):
            doc.add_paragraph()
            i += 1
            continue

        m = re.match(r"^\s*[-*]\s+(.+)$", line)
        if m:
            add_styled_paragraph(m.group(1).strip(), bullet_style)
            i += 1
            continue

        m = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if m:
            add_styled_paragraph(m.group(1).strip(), number_style)
            i += 1
            continue

        if stripped == "":
            i += 1
            continue

        # Default: paragraph with full inline formatting (bold, italic, code)
        p = doc.add_paragraph()
        _render_inline(p, stripped)
        i += 1


def _move_elements_after(body, anchor_el, elements: list) -> None:
    """Move XML elements (in given order) to sit directly after anchor_el."""
    cur = anchor_el
    for el in elements:
        if el.getparent() is body:
            body.remove(el)
        cur.addnext(el)
        cur = el


def _insert_markdown_after(doc, anchor_paragraph, content_text: str) -> None:
    """Render content_text (markdown-ish) as new paragraphs/tables and place
    them directly after anchor_paragraph in the document body."""
    body = doc.element.body

    # Snapshot existing paragraph+table elements
    before = set()
    for p in doc.paragraphs:
        before.add(p._element)
    for t in doc.tables:
        before.add(t._element)

    if content_text.strip() == MISSING_MARKER:
        p = doc.add_paragraph()
        run = p.add_run("⚠ Keine Informationen aus den hochgeladenen Dokumenten verfügbar.")
        run.italic = True
    else:
        _append_markdown_to_doc(doc, content_text)

    # Collect new elements
    new_elements = []
    for p in doc.paragraphs:
        if p._element not in before:
            new_elements.append(p._element)
    for t in doc.tables:
        if t._element not in before:
            new_elements.append(t._element)

    # Preserve document order
    order = {el: idx for idx, el in enumerate(body.iterchildren())}
    new_elements.sort(key=lambda el: order.get(el, 0))

    _move_elements_after(body, anchor_paragraph._element, new_elements)


def _fill_template_with_sections(doc, sections_lookup: dict) -> set:
    """Walk every Heading paragraph in the template, replace its placeholder
    body (everything between this heading and the next heading) with the
    matching agent section, or with a 'no info' marker if missing.

    Returns the set of normalized heading keys that were consumed.
    """
    heading_paragraphs = [
        p for p in doc.paragraphs if p.style.name.startswith("Heading ")
    ]
    if not heading_paragraphs:
        return set()

    body = doc.element.body
    consumed = set()

    # Process from LAST heading to FIRST so removing/inserting elements
    # earlier in the body doesn't shift the targets we still need to find.
    for idx in range(len(heading_paragraphs) - 1, -1, -1):
        h = heading_paragraphs[idx]
        h_key = _normalize_heading(h.text)

        next_h_el = (
            heading_paragraphs[idx + 1]._element
            if idx + 1 < len(heading_paragraphs)
            else None
        )

        # Collect all body elements between this heading and the next
        # (excluding sectPr at the end).
        to_remove = []
        cur = h._element.getnext()
        while cur is not None and cur is not next_h_el:
            nxt = cur.getnext()
            tag = cur.tag.rsplit("}", 1)[-1] if "}" in cur.tag else cur.tag
            if tag == "sectPr":
                break
            to_remove.append(cur)
            cur = nxt

        for el in to_remove:
            if el.getparent() is body:
                body.remove(el)

        content = sections_lookup.get(h_key)
        content_stripped = (content or "").strip()
        # Only count as actually filled when the agent produced real content,
        # not when it emitted the "no info available" marker.
        if content_stripped and content_stripped != MISSING_MARKER:
            consumed.add(h_key)
        if not content_stripped:
            content = MISSING_MARKER

        _insert_markdown_after(doc, h, content)

    return consumed


@app.post("/api/generate-isdp-docx")
async def generate_isdp_docx(req: DraftRequest):
    if not os.path.exists(TEMPLATE_PATH):
        raise HTTPException(500, f"Template not found at {TEMPLATE_PATH}")
    if not req.draft_text.strip():
        raise HTTPException(400, "draft_text is empty")

    from docx import Document

    doc = Document(TEMPLATE_PATH)
    project = (req.project_name or "ISDP").strip()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    sections = _parse_agent_sections(req.draft_text)
    logger.info(
        "Parsed %d sections from agent draft (input length: %d chars)",
        len(sections), len(req.draft_text),
    )
    if len(sections) == 0:
        preview = req.draft_text[:400].replace("\n", " ⏎ ")
        logger.warning(
            "No sections parsed. Draft input preview: %r", preview
        )
    else:
        logger.info(
            "Matched template chapters: %s",
            ", ".join(sorted(sections.keys()))[:500],
        )

    consumed = _fill_template_with_sections(doc, sections)

    # Append any agent sections that didn't match a template heading
    unmatched = {k: v for k, v in sections.items() if k not in consumed and v.strip()}
    if unmatched:
        doc.add_page_break()
        doc.add_heading("Zusätzliche Informationen (nicht zugeordnet)", level=1)
        doc.add_paragraph(
            "Diese Abschnitte wurden vom Agenten erstellt, konnten aber keinem "
            "Vorlage-Kapitel zugeordnet werden."
        )
        for key, content in unmatched.items():
            doc.add_heading(key.title(), level=2)
            _append_markdown_to_doc(doc, content)

    # Footer note about generation
    doc.add_paragraph()
    total = len(ISDS_SECTIONS)
    filled = len(consumed)
    missing = total - filled
    missing_note = f" · {missing} Kapitel ohne verfügbare Quellinformation" if missing else ""
    note = doc.add_paragraph()
    run = note.add_run(
        f"Automatisch erstellt am {datetime.now().strftime('%d.%m.%Y %H:%M')} "
        f"durch den ISDP-Assistenten (Projekt: {project}). "
        f"Kapitel mit Inhalt: {filled} / {total}{missing_note}."
    )
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_project = re.sub(r"[^A-Za-z0-9_-]+", "_", project).strip("_") or "ISDP"
    filename = f"ISDP_Draft_{safe_project}_{timestamp}.docx"
    logger.info(
        "Generated ISDP draft docx: %s (%d sections parsed, %d filled, %d unmatched)",
        filename, len(sections), len(consumed), len(unmatched),
    )

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "X-Sections-Parsed": str(len(sections)),
            "X-Sections-Filled": str(len(consumed)),
            "X-Sections-Total": str(len(ISDS_SECTIONS)),
        },
    )


# ── Serve frontend (must be last) ─────────────────────────────────────────────
app.mount("/", StaticFiles(directory="public", html=True), name="static")